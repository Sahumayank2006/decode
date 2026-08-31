"""
DECODE – OCR Engine
Combines OpenCV image pre-processing with Tesseract OCR.
Supports: single images, multi-page PDFs, DOCX files.
"""

import os
import io
import logging
import re
from pathlib import Path
from typing import Optional

import cv2
import numpy as np
from PIL import Image

# EasyOCR will be loaded lazily to prevent reloading models for every request
logger = logging.getLogger("decode.ocr")
ocr_reader = None

def get_ocr_reader():
    global ocr_reader
    if ocr_reader is None:
        logger.info("Initializing EasyOCR models (this may take a moment on first run)...")
        try:
            import easyocr
            ocr_reader = easyocr.Reader(['en']) # Will use GPU if available, else CPU
        except Exception as e:
            logger.error(f"Failed to initialize EasyOCR: {e}")
    return ocr_reader

# ─────────────────────────────────────────────────────────────────────────────
# Image pre-processing helpers
# ─────────────────────────────────────────────────────────────────────────────

def _load_image(source) -> np.ndarray:
    """Accept file path, bytes, PIL image, or numpy array."""
    if isinstance(source, np.ndarray):
        return source
    if isinstance(source, (str, Path)):
        img = cv2.imread(str(source))
        if img is None:
            raise ValueError(f"Cannot load image: {source}")
        return img
    if isinstance(source, bytes):
        arr = np.frombuffer(source, dtype=np.uint8)
        return cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if isinstance(source, Image.Image):
        return cv2.cvtColor(np.array(source), cv2.COLOR_RGB2BGR)
    raise TypeError(f"Unsupported image source type: {type(source)}")


def preprocess_image(img: np.ndarray,
                     deskew: bool = True,
                     denoise: bool = True,
                     enhance_contrast: bool = True,
                     binarize: bool = True) -> np.ndarray:
    """
    Full OpenCV pre-processing pipeline for OCR-optimized images.
    Steps: resize → grayscale → denoise → contrast → binarize → deskew
    """
    # 1. Resize to fixed width if too small
    h, w = img.shape[:2]
    if w < 800:
        scale = 1200 / w
        img = cv2.resize(img, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)

    # 2. Grayscale
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img.copy()

    # 3. Denoise
    if denoise:
        gray = cv2.fastNlMeansDenoising(gray, h=10, templateWindowSize=7, searchWindowSize=21)

    # 4. Contrast enhancement (CLAHE)
    if enhance_contrast:
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        gray = clahe.apply(gray)

    # 5. Binarize (adaptive threshold)
    if binarize:
        gray = cv2.adaptiveThreshold(
            gray, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, 11, 2
        )

    # 6. Deskew
    if deskew:
        gray = _deskew(gray)

    return gray


def _deskew(img: np.ndarray) -> np.ndarray:
    """Correct skew using Hough transform."""
    try:
        coords = np.column_stack(np.where(img > 0))
        if len(coords) < 50:
            return img
        angle = cv2.minAreaRect(coords.astype(np.float32))[-1]
        if angle < -45:
            angle = -(90 + angle)
        else:
            angle = -angle
        if abs(angle) < 0.5:          # negligible skew
            return img
        h, w = img.shape[:2]
        M = cv2.getRotationMatrix2D((w // 2, h // 2), angle, 1.0)
        return cv2.warpAffine(img, M, (w, h),
                              flags=cv2.INTER_CUBIC,
                              borderMode=cv2.BORDER_REPLICATE)
    except Exception as e:
        logger.debug("Deskew skipped: %s", e)
        return img


def remove_noise(img: np.ndarray) -> np.ndarray:
    """Remove small noise blobs via morphological opening."""
    kernel = np.ones((1, 1), np.uint8)
    img = cv2.dilate(img, kernel, iterations=1)
    img = cv2.erode(img, kernel, iterations=1)
    img = cv2.morphologyEx(img, cv2.MORPH_CLOSE, kernel)
    return img


def detect_tables(img: np.ndarray) -> list[dict]:
    """
    Detect table regions in the image using line detection.
    Returns a list of bounding-box dicts: {x, y, w, h}.
    """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img
    binary = cv2.adaptiveThreshold(gray, 255,
                                   cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                   cv2.THRESH_BINARY_INV, 15, 2)

    # Horizontal lines
    h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
    h_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, h_kernel)

    # Vertical lines
    v_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
    v_lines = cv2.morphologyEx(binary, cv2.MORPH_OPEN, v_kernel)

    # Combine
    table_mask = cv2.addWeighted(h_lines, 0.5, v_lines, 0.5, 0.0)
    _, table_bin = cv2.threshold(table_mask, 50, 255, cv2.THRESH_BINARY)

    contours, _ = cv2.findContours(table_bin, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    tables = []
    for cnt in contours:
        area = cv2.contourArea(cnt)
        if area > 5000:
            x, y, w, h = cv2.boundingRect(cnt)
            tables.append({"x": int(x), "y": int(y), "w": int(w), "h": int(h), "area": int(area)})

    return sorted(tables, key=lambda t: t["area"], reverse=True)


def detect_figures(img: np.ndarray) -> list[dict]:
    """
    Detect figure/image regions in a document image.
    Returns bounding boxes for non-text regions.
    """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img
    _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY_INV)

    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (15, 15))
    closed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

    contours, _ = cv2.findContours(closed, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    figures = []
    for cnt in contours:
        area = cv2.contourArea(cnt)
        if area > 10000:
            x, y, w, h = cv2.boundingRect(cnt)
            aspect = w / max(h, 1)
            if 0.2 < aspect < 5.0:
                figures.append({
                    "x": int(x), "y": int(y),
                    "w": int(w), "h": int(h),
                    "area": int(area),
                    "aspect_ratio": round(aspect, 2)
                })
    return figures


# ─────────────────────────────────────────────────────────────────────────────
# OCR core
# ─────────────────────────────────────────────────────────────────────────────

def ocr_image(source,
              lang: str = "eng",
              preprocess: bool = True) -> dict:
    """
    Run EasyOCR on a single image.
    Returns: {text, confidence, word_count, char_count, lines}
    """
    img = _load_image(source)
    raw_img = img.copy()

    if preprocess:
        processed = preprocess_image(img)
    else:
        processed = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY) if len(img.shape) == 3 else img

    reader = get_ocr_reader()
    if not reader:
        raise RuntimeError("EasyOCR reader is not initialized.")

    try:
        # EasyOCR readtext returns a list of tuples: (bbox, text, prob)
        results = reader.readtext(processed)
        
        words = []
        confs = []
        lines = []
        
        for (bbox, text, prob) in results:
            if text.strip():
                lines.append(text.strip())
                # Split text into words for count
                words.extend(text.split())
                confs.append(prob * 100) # EasyOCR prob is 0.0 to 1.0
                
        full_text = "\n".join(lines)
        avg_conf = round(sum(confs) / len(confs), 2) if confs else 0.0

        return {
            "text": full_text.strip(),
            "lines": lines,
            "word_count": len(words),
            "char_count": len(full_text.strip()),
            "confidence": avg_conf,
            "tables_detected": detect_tables(raw_img),
            "figures_detected": detect_figures(raw_img),
            "language": lang,
        }
    except Exception as e:
        logger.error(f"OCR Error: {e}")
        raise


def ocr_pdf(pdf_path: str,
            lang: str = "eng",
            dpi: int = 300) -> dict:
    """
    OCR a multi-page PDF.
    Converts each page to image then runs OCR.
    Returns per-page results + merged text.
    """
    try:
        from pdf2image import convert_from_path
    except ImportError:
        raise RuntimeError("pdf2image is required for PDF OCR. Install it via pip.")

    logger.info("OCR PDF: %s", pdf_path)
    pages = convert_from_path(pdf_path, dpi=dpi)

    page_results = []
    for i, page_img in enumerate(pages):
        logger.debug("  Processing page %d/%d", i + 1, len(pages))
        result = ocr_image(np.array(page_img), lang=lang)
        result["page"] = i + 1
        page_results.append(result)

    merged_text = "\n\n".join(r["text"] for r in page_results)
    total_words = sum(r["word_count"] for r in page_results)
    avg_conf = round(
        sum(r["confidence"] for r in page_results) / len(page_results), 2
    ) if page_results else 0.0

    return {
        "text": merged_text,
        "pages": page_results,
        "total_pages": len(pages),
        "total_words": total_words,
        "average_confidence": avg_conf,
        "source": str(pdf_path),
    }


def ocr_docx(docx_path: str) -> dict:
    """
    Extract text from DOCX without OCR (native text extraction).
    Falls back to OCR for embedded images.
    """
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx is required.")

    doc = Document(docx_path)
    paragraphs, tables_text = [], []

    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(" | ".join(cells))
        tables_text.extend(table_rows)

    full_text = "\n".join(paragraphs)
    table_text = "\n".join(tables_text)
    combined = full_text + ("\n\nTABLES:\n" + table_text if table_text else "")

    return {
        "text": combined,
        "paragraphs": paragraphs,
        "tables": tables_text,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "confidence": 100.0,   # native extraction = perfect
        "source": str(docx_path),
    }


def extract_text_from_file(file_path: str, lang: str = "eng") -> dict:
    """
    Dispatcher: choose OCR method based on file extension.
    """
    ext = Path(file_path).suffix.lower()
    logger.info("Extracting text from: %s (ext=%s)", file_path, ext)

    if ext == ".pdf":
        return ocr_pdf(file_path, lang=lang)
    elif ext in (".docx", ".doc"):
        return ocr_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif", ".webp"):
        return ocr_image(file_path, lang=lang)
    elif ext == ".txt":
        with open(file_path, "r", errors="replace") as f:
            text = f.read()
        return {
            "text": text,
            "word_count": len(text.split()),
            "char_count": len(text),
            "confidence": 100.0,
        }
    else:
        raise ValueError(f"Unsupported file type: {ext}")
